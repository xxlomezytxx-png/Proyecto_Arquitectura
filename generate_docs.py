"""
Genera documentación Word para cada DEV del Sprint 2 de BookFlow AI Commerce.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ────────────────────────────────────────────────────────────
# Helpers de estilo
# ────────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1A, 0x37, 0x6E)   # encabezados principales
AZUL_MEDIO   = RGBColor(0x1F, 0x6F, 0xB8)   # h2
GRIS_TEXTO   = RGBColor(0x33, 0x33, 0x33)
VERDE        = RGBColor(0x1A, 0x7A, 0x3C)
NARANJA      = RGBColor(0xB8, 0x5C, 0x00)
GRIS_FONDO   = RGBColor(0xF2, 0xF2, 0xF2)   # fondo de código
NEGRO        = RGBColor(0x00, 0x00, 0x00)

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code(doc, text):
    """Bloque de código con fuente monoespaciada y fondo gris."""
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        # fondo gris en el párrafo (shading)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
    doc.add_paragraph()  # espacio después del bloque

def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Encabezado
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, "1A376E")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(10)

    # Filas de datos
    for r_idx, row in enumerate(rows):
        data_row = table.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "EEF4FB"
        for c_idx, cell_text in enumerate(row):
            cell = data_row.cells[c_idx]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A376E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def cover_page(doc):
    doc.add_picture  # no usamos imagen
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("BookFlow AI Commerce")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = AZUL_OSCURO
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Documentación Técnica — Sprint 2")
    run2.font.size  = Pt(18)
    run2.font.color.rgb = AZUL_MEDIO
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = doc.add_paragraph()
    run3 = p3.add_run("Motor de IA, Pricing Trazable e Integración Real")
    run3.font.size  = Pt(13)
    run3.font.color.rgb = GRIS_TEXTO
    run3.italic = True
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    p4 = doc.add_paragraph()
    run4 = p4.add_run("Guía por Desarrollador: archivos, endpoints y pruebas")
    run4.font.size  = Pt(11)
    run4.font.color.rgb = GRIS_TEXTO
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p5 = doc.add_paragraph()
    run5 = p5.add_run("Versión 1.0  ·  Abril 2026")
    run5.font.size  = Pt(10)
    run5.font.color.rgb = GRIS_TEXTO
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# SECCIÓN ÍNDICE
# ────────────────────────────────────────────────────────────
def index_page(doc):
    add_heading(doc, "Índice de Desarrolladores", 1, AZUL_OSCURO)
    devs = [
        ("DEV 1", "AI Enrichment Service — Servicio Real",           "Fuentes externas: Google Books, Open Library, Crossref"),
        ("DEV 2", "Limpieza y Normalización de Datos",               "Normalizer: ISBN, autores, títulos, merge multi-fuente"),
        ("DEV 3", "Pricing Service — Motor de Precios",              "eBay, fallback interno, factores de condición, trazabilidad"),
        ("DEV 4", "UI de Pricing — Panel Administrativo",            "React: PricingDashboard, PricingCard, PricingHistory"),
        ("DEV 5", "Catálogo Enriquecido — Frontend Comercial",       "React: CatalogPage, BookCard, BookDetail, filtros, búsqueda"),
        ("DEV 6", "Integración Real entre Microservicios",           "BFF Gateway: rutas Sprint 2, health check consolidado"),
        ("DEV 7", "Auditoría de Pricing e IA (BFF)",                 "Audit endpoints, anomalías, estadísticas"),
        ("DEV 8", "Integración con APIs Externas — Robustez",        "Circuit breaker, caché en memoria, estado de APIs"),
        ("DEV 9", "Optimización y Persistencia",                     "docker-compose, BD, seeds, índices"),
    ]
    add_table(doc,
        ["DEV", "Responsable de", "Descripción rápida"],
        devs,
        [0.6, 2.4, 3.0]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 1 — AI Enrichment Service
# ────────────────────────────────────────────────────────────
def dev1(doc):
    add_heading(doc, "DEV 1 — AI Enrichment Service (Servicio Real)", 1, AZUL_OSCURO)
    add_body(doc,
        "Responsable de convertir el mock de enriquecimiento del Sprint 1 en un servicio real "
        "que consulta APIs externas (Google Books, Open Library, Crossref), normaliza los datos "
        "y los persiste en enrichment_db.", color=GRIS_TEXTO)
    add_separator(doc)

    # Archivos
    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("ai-enrichment-service/app/main.py",                           "Modificado", "Lifespan, CORS, montaje de router"),
            ("ai-enrichment-service/app/config.py",                          "Nuevo",      "Settings con Pydantic (URL de servicios, API keys)"),
            ("ai-enrichment-service/app/domain/entities/enrichment.py",      "Nuevo",      "Entidades EnrichmentRequest, EnrichmentResult, BookMetadata"),
            ("ai-enrichment-service/app/domain/interfaces/enrichment_repository.py", "Nuevo", "Contrato abstracto del repositorio"),
            ("ai-enrichment-service/app/infrastructure/database/models.py",  "Nuevo",      "Modelos SQLAlchemy: EnrichmentRequestModel, EnrichmentResultModel"),
            ("ai-enrichment-service/app/infrastructure/database/connection.py","Nuevo",    "Sesión async de SQLAlchemy + creación de tablas"),
            ("ai-enrichment-service/app/infrastructure/database/repositories/enrichment_repository.py","Nuevo","Implementación del repositorio"),
            ("ai-enrichment-service/app/infrastructure/adapters/base_adapter.py","Nuevo",  "SimpleCache TTL + circuit breaker reutilizable"),
            ("ai-enrichment-service/app/infrastructure/adapters/google_books_adapter.py","Nuevo","Adapter Google Books API"),
            ("ai-enrichment-service/app/infrastructure/adapters/open_library_adapter.py","Nuevo","Adapter Open Library API"),
            ("ai-enrichment-service/app/infrastructure/adapters/crossref_adapter.py","Nuevo","Adapter Crossref API"),
            ("ai-enrichment-service/app/application/use_cases/enrich_book.py","Nuevo",     "Caso de uso: orquesta adapters + normalizer + repositorio"),
            ("ai-enrichment-service/app/routers/enrichment_router.py",       "Nuevo",      "Endpoints FastAPI: /enrich, /enrich/{id}, /external-apis/status, /health"),
            ("ai-enrichment-service/tests/unit/test_enrich_use_case.py",     "Nuevo",      "Tests unitarios del caso de uso"),
            ("ai-enrichment-service/requirements.txt",                        "Modificado", "Dependencias: fastapi, sqlalchemy, httpx, pydantic-settings"),
            ("ai-enrichment-service/README.md",                               "Nuevo",      "Documentación completa del servicio"),
        ],
        [3.5, 1.0, 3.5]
    )

    # Endpoints
    add_heading(doc, "Endpoints disponibles", 2, AZUL_MEDIO)
    add_table(doc,
        ["Método", "Ruta", "Descripción"],
        [
            ("POST", "/enrich",                      "Enriquecer un libro (por ISBN o título/autor)"),
            ("GET",  "/enrich/{request_id}",         "Consultar estado y resultado de una solicitud"),
            ("GET",  "/external-apis/status",        "Estado de circuit breaker de cada API externa"),
            ("GET",  "/health",                      "Health check del servicio"),
        ],
        [0.8, 2.2, 4.0]
    )

    # Cómo probar
    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Levantar todos los servicios:", bold=True)
    add_code(doc, "docker compose up --build")

    add_body(doc, "2. Enriquecer un libro por ISBN:", bold=True)
    add_code(doc, """\
curl -X POST http://localhost:8004/enrich \\
  -H "Content-Type: application/json" \\
  -d '{
    "book_id": "550e8400-e29b-41d4-a716-446655440001",
    "isbn": "9780141439518",
    "title": "Great Expectations",
    "author": "Charles Dickens"
  }'""")

    add_body(doc, "Respuesta esperada:", italic=True)
    add_code(doc, """\
{
  "request_id": "7b4e2a1c-...",
  "status": "completed",
  "normalized_title": "Great Expectations",
  "normalized_author": "Dickens, Charles",
  "cover_url": "https://books.google.com/...",
  "source_used": "google_books",
  "confidence_score": 0.95
}""")

    add_body(doc, "3. Consultar el estado de una solicitud:", bold=True)
    add_code(doc, "curl http://localhost:8004/enrich/{request_id}")

    add_body(doc, "4. Ver estado de las APIs externas:", bold=True)
    add_code(doc, "curl http://localhost:8004/external-apis/status")

    add_body(doc, "5. Ejecutar los tests unitarios:", bold=True)
    add_code(doc, """\
cd ai-enrichment-service
pytest tests/ -v --cov=app --cov-report=term-missing""")

    # Tests
    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Archivo", "Qué verifica"],
        [
            ("test_enrich_by_isbn_google_books_success",      "test_enrich_use_case.py", "Google Books devuelve datos correctos"),
            ("test_enrich_fallback_when_google_fails",        "test_enrich_use_case.py", "Fallback a Open Library si Google falla"),
            ("test_enrich_all_sources_fail_returns_partial",  "test_enrich_use_case.py", "Datos parciales si todas las fuentes fallan"),
            ("test_normalization_merges_fields_correctly",    "test_enrich_use_case.py", "Normalizer fusiona campos multi-fuente"),
            ("test_adapter_timeout_handled_gracefully",       "test_enrich_use_case.py", "Timeout no levanta excepción no manejada"),
        ],
        [2.8, 2.0, 3.2]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 2 — Normalizer
# ────────────────────────────────────────────────────────────
def dev2(doc):
    add_heading(doc, "DEV 2 — Limpieza y Normalización de Datos", 1, AZUL_OSCURO)
    add_body(doc,
        "Módulo dentro de ai-enrichment-service que homologa datos provenientes de fuentes "
        "heterogéneas: normaliza títulos, autores, valida ISBNs, fusiona resultados y detecta "
        "duplicados antes de persistir.", color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("ai-enrichment-service/app/application/normalizer/__init__.py",    "Nuevo", "Exporta el paquete"),
            ("ai-enrichment-service/app/application/normalizer/text_normalizer.py",  "Nuevo", "strip, title_case, elimina caracteres de control"),
            ("ai-enrichment-service/app/application/normalizer/isbn_validator.py",   "Nuevo", "Validación checksum ISBN-10/13, conversión ISBN-10→13"),
            ("ai-enrichment-service/app/application/normalizer/author_formatter.py", "Nuevo", "Convierte 'Nombre Apellido' → 'Apellido, Nombre'; múltiples autores con ';'"),
            ("ai-enrichment-service/app/application/normalizer/source_merger.py",    "Nuevo", "Fusiona resultados por confidence_score; prefiere fuente con mayor score"),
            ("ai-enrichment-service/tests/unit/test_normalizer.py",                  "Nuevo", "Tests unitarios de todos los módulos del normalizer"),
        ],
        [3.5, 1.0, 3.5]
    )

    add_heading(doc, "Reglas de normalización implementadas", 2, AZUL_MEDIO)
    add_table(doc,
        ["Campo", "Regla"],
        [
            ("Título",   "strip() + title_case + eliminar caracteres de control (\\r \\n \\t)"),
            ("Autor",    "'Nombre Apellido' → 'Apellido, Nombre'; múltiples autores separados por ';'"),
            ("ISBN-10",  "Validar checksum módulo 11 y convertir a ISBN-13"),
            ("ISBN-13",  "Validar checksum módulo 10"),
            ("Conflicto entre fuentes", "Se prefiere la fuente con confidence_score más alto"),
        ],
        [2.0, 6.0]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "Los módulos del normalizer se testean de forma aislada (sin Docker):", bold=True)
    add_code(doc, """\
cd ai-enrichment-service
pip install -r requirements.txt
pytest tests/unit/test_normalizer.py -v""")

    add_body(doc, "Ejemplo de uso directo en Python:", italic=True)
    add_code(doc, """\
from app.application.normalizer.isbn_validator import validate_isbn, isbn10_to_isbn13
from app.application.normalizer.author_formatter import format_author

print(isbn10_to_isbn13("0141439513"))   # → "9780141439518"
print(format_author("Charles Dickens")) # → "Dickens, Charles"
print(format_author("Gabriel García Márquez")) # → "García Márquez, Gabriel"
""")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Qué verifica"],
        [
            ("test_isbn10_converts_to_isbn13_correctly",      "Conversión ISBN-10 → ISBN-13 produce dígito verificador correcto"),
            ("test_invalid_isbn_raises_validation_error",     "ISBN con checksum incorrecto lanza error"),
            ("test_author_format_single_name",                "Nombre simple se convierte a 'Apellido, Nombre'"),
            ("test_author_format_multiple_authors",           "Múltiples autores quedan separados por ';'"),
            ("test_title_normalization_removes_special_chars","Título limpio sin caracteres de control"),
            ("test_source_merger_prefers_higher_confidence",  "source_merger selecciona el campo con mayor confidence_score"),
            ("test_duplicate_detection_by_isbn",              "ISBN duplicado es detectado correctamente"),
        ],
        [3.2, 4.8]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 3 — Pricing Service
# ────────────────────────────────────────────────────────────
def dev3(doc):
    add_heading(doc, "DEV 3 — Pricing Service (Motor de Precios)", 1, AZUL_OSCURO)
    add_body(doc,
        "Microservicio independiente que calcula precios sugeridos trazables. Consulta eBay Browse API "
        "para obtener referencias de mercado, aplica factores de condición física y persiste cada "
        "decisión con explicación en español. Si eBay no responde, aplica reglas internas como fallback.",
        color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("pricing-service/app/main.py",                                     "Nuevo",      "Lifespan, CORS, montaje de router, health check"),
            ("pricing-service/app/config.py",                                    "Nuevo",      "Settings: DATABASE_URL, EBAY_APP_ID, EBAY_BASE_URL, SERVICE_PORT"),
            ("pricing-service/app/domain/entities/pricing.py",                   "Nuevo",      "Entidades: PriceReference, PricingDecision, ConditionFactor (enum)"),
            ("pricing-service/app/domain/interfaces/pricing_repository.py",      "Nuevo",      "Contrato abstracto: save_decision, get_by_book, get_history"),
            ("pricing-service/app/infrastructure/database/models.py",            "Nuevo",      "Modelos SQLAlchemy: PriceReferenceModel, PricingDecisionModel"),
            ("pricing-service/app/infrastructure/database/connection.py",         "Nuevo",      "Sesión async SQLAlchemy + creación de tablas"),
            ("pricing-service/app/infrastructure/database/repositories/pricing_repository.py","Nuevo","Implementación CRUD del repositorio"),
            ("pricing-service/app/infrastructure/adapters/ebay_adapter.py",      "Nuevo",      "Adapter eBay Browse API con circuit breaker y caché"),
            ("pricing-service/app/application/use_cases/calculate_price.py",     "Nuevo",      "Caso de uso: referencia eBay + fallback interno + persistencia"),
            ("pricing-service/app/routers/pricing_router.py",                     "Nuevo",      "Endpoints FastAPI: calculate, detalle, historial, explicación, APIs status"),
            ("pricing-service/tests/unit/test_calculate_price.py",               "Nuevo",      "Tests unitarios del motor de pricing"),
            ("pricing-service/requirements.txt",                                  "Nuevo",      "Dependencias del servicio"),
            ("pricing-service/README.md",                                          "Nuevo",      "Documentación completa"),
        ],
        [3.5, 1.0, 3.5]
    )

    add_heading(doc, "Endpoints disponibles", 2, AZUL_MEDIO)
    add_table(doc,
        ["Método", "Ruta", "Descripción"],
        [
            ("POST", "/pricing/calculate",               "Calcular precio sugerido para un libro"),
            ("GET",  "/pricing/{book_id}",               "Último precio calculado para un libro"),
            ("GET",  "/pricing/{book_id}/history",       "Historial de precios de un libro"),
            ("GET",  "/pricing/{book_id}/explanation",   "Desglose legible del cálculo"),
            ("GET",  "/pricing/external-apis/status",    "Estado del circuit breaker de eBay"),
            ("GET",  "/health",                          "Health check"),
        ],
        [0.8, 2.5, 4.7]
    )

    add_heading(doc, "Factores de condición", 2, AZUL_MEDIO)
    add_table(doc,
        ["Condición", "Factor", "Efecto sobre el precio base"],
        [
            ("NUEVO",       "1.00", "Precio de referencia completo"),
            ("BUENO",       "0.75", "75 % del precio de referencia"),
            ("ACEPTABLE",   "0.50", "50 % del precio de referencia"),
            ("DETERIORADO", "0.25", "25 % del precio de referencia"),
        ],
        [2.0, 1.2, 4.8]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Levantar todos los servicios:", bold=True)
    add_code(doc, "docker compose up --build")

    add_body(doc, "2. Calcular precio de un libro:", bold=True)
    add_code(doc, """\
curl -X POST http://localhost:8005/pricing/calculate \\
  -H "Content-Type: application/json" \\
  -d '{
    "book_id": "550e8400-e29b-41d4-a716-446655440001",
    "isbn": "9780141439518",
    "title": "Great Expectations",
    "condition": "BUENO",
    "publication_year": 2010
  }'""")

    add_body(doc, "Respuesta esperada:", italic=True)
    add_code(doc, """\
{
  "book_id": "550e8400-...",
  "suggested_price": 9.38,
  "currency": "USD",
  "explanation": "Precio calculado con base en 3 referencias de eBay (promedio: $12.50). Factor de condición BUENO aplicado (0.75). Precio sugerido: $9.38.",
  "source": "ebay",
  "condition_factor": 0.75,
  "is_fallback": false
}""")

    add_body(doc, "3. Ver historial de precios de un libro:", bold=True)
    add_code(doc, "curl http://localhost:8005/pricing/{book_id}/history")

    add_body(doc, "4. Ver desglose del cálculo:", bold=True)
    add_code(doc, "curl http://localhost:8005/pricing/{book_id}/explanation")

    add_body(doc, "5. Ejecutar los tests:", bold=True)
    add_code(doc, """\
cd pricing-service
pytest tests/ -v --cov=app --cov-report=term-missing""")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Qué verifica"],
        [
            ("test_calculate_price_new_book_with_references",         "Libro NUEVO usa referencias eBay sin reducción de factor"),
            ("test_calculate_price_used_book_condition_factor_applied","Factor de condición se multiplica correctamente sobre precio base"),
            ("test_fallback_when_ebay_unavailable",                    "Fallback interno produce precio cuando eBay no responde"),
            ("test_minimum_price_never_below_threshold",               "Precio mínimo absoluto de $1.00 siempre respetado"),
            ("test_explanation_text_generated_correctly",              "Campo explanation contiene fuente, factor y precio final"),
            ("test_pricing_decision_persisted_with_full_traceability", "PricingDecision se guarda con todos sus campos"),
            ("test_condition_factor_deteriorated_applies_correctly",   "Factor 0.25 aplicado correctamente en condición DETERIORADO"),
        ],
        [3.5, 4.5]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 4 — UI de Pricing
# ────────────────────────────────────────────────────────────
def dev4(doc):
    add_heading(doc, "DEV 4 — UI de Pricing (Panel Administrativo)", 1, AZUL_OSCURO)
    add_body(doc,
        "Componentes React en admin-frontend para visualizar y gestionar precios sugeridos. "
        "Muestra desglose del cálculo, historial de cambios, badge de fuente (Verificado/Estimado) "
        "y botón de recálculo. Se comunica exclusivamente con el BFF.", color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("admin-frontend/src/components/pricing/PricingDashboard.jsx", "Nuevo",      "Vista principal: tabla de libros con precio sugerido"),
            ("admin-frontend/src/components/pricing/PricingCard.jsx",      "Nuevo",      "Tarjeta de detalle de un precio con badge de fuente"),
            ("admin-frontend/src/components/pricing/PricingExplanation.jsx","Nuevo",     "Desglose visual del cálculo (breakdown)"),
            ("admin-frontend/src/components/pricing/PricingHistory.jsx",   "Nuevo",      "Historial de precios de un libro"),
            ("admin-frontend/src/components/pricing/RecalculateButton.jsx","Nuevo",      "Botón de recálculo con loading state"),
            ("admin-frontend/src/components/pricing/PricingFilters.jsx",   "Nuevo",      "Filtros: condición, rango de precio, estado"),
            ("admin-frontend/src/pages/PreciosPage.jsx",                   "Nuevo",      "Página que monta PricingDashboard + filtros"),
            ("admin-frontend/src/services/pricingService.js",             "Nuevo",      "Llamadas al BFF: lista, detalle, recalcular, historial"),
            ("admin-frontend/src/App.js",                                  "Modificado", "Añade ruta /precios → PreciosPage"),
            ("admin-frontend/src/components/Sidebar.jsx",                  "Nuevo",      "Menú lateral con navegación hacia todas las páginas"),
            ("admin-frontend/src/services/api.js",                         "Nuevo",      "Instancia axios con baseURL apuntando al BFF"),
        ],
        [3.5, 1.0, 3.5]
    )

    add_heading(doc, "Rutas del BFF que consume", 2, AZUL_MEDIO)
    add_table(doc,
        ["Método", "Ruta en BFF", "Para qué"],
        [
            ("GET",  "/api/admin/pricing",                "Listado de precios de todos los libros"),
            ("GET",  "/api/admin/pricing/{book_id}",      "Detalle con explicación"),
            ("POST", "/api/admin/pricing/recalculate",    "Solicitar recálculo de precio"),
            ("GET",  "/api/admin/pricing/{id}/history",   "Historial de precios"),
        ],
        [0.8, 2.8, 4.4]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Levantar el proyecto completo:", bold=True)
    add_code(doc, "docker compose up --build")

    add_body(doc, "2. Abrir el panel admin en el navegador:", bold=True)
    add_code(doc, "http://localhost:3001")

    add_body(doc, "3. Navegar a la sección Precios en el menú lateral.", bold=False)
    add_body(doc, "4. Verificar que la tabla carga los libros con precios sugeridos.", bold=False)
    add_body(doc, "5. Hacer clic en un libro → ver el desglose del cálculo en PricingExplanation.", bold=False)
    add_body(doc, "6. Verificar badges: verde 'Verificado' si source = ebay, naranja 'Estimado' si is_fallback = true.", bold=False)
    add_body(doc, "7. Usar RecalculateButton y verificar estado de carga.", bold=False)

    add_body(doc, "8. Desarrollo local sin Docker:", bold=True)
    add_code(doc, """\
cd admin-frontend
npm install
npm start
# → http://localhost:3001""")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Componente", "Qué verifica"],
        [
            ("PricingDashboard renders table with pricing data", "PricingDashboard", "Tabla se renderiza con datos de precios"),
            ("PricingCard shows fallback badge when source is internal","PricingCard","Badge 'Estimado' en naranja cuando is_fallback=true"),
            ("RecalculateButton shows loading state during request",    "RecalculateButton","Spinner visible durante la petición"),
            ("PricingExplanation displays breakdown correctly",         "PricingExplanation","Campos de desglose visibles en la UI"),
            ("Error state renders user-friendly message",              "PricingDashboard","Mensaje legible al recibir error del BFF"),
        ],
        [3.0, 1.5, 3.5]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 5 — Catálogo Enriquecido
# ────────────────────────────────────────────────────────────
def dev5(doc):
    add_heading(doc, "DEV 5 — Catálogo Enriquecido (Frontend Comercial)", 1, AZUL_OSCURO)
    add_body(doc,
        "Actualización de commercial-frontend para mostrar productos con datos enriquecidos "
        "(portada, descripción, precio integrado), filtros avanzados y búsqueda con debounce. "
        "Todo se comunica exclusivamente con el BFF.", color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("commercial-frontend/src/pages/Catalogo.jsx",              "Nuevo",      "Página principal del catálogo: grid de libros + paginación"),
            ("commercial-frontend/src/components/BookCard.jsx",          "Modificado", "Tarjeta de libro con precio, condición y PriceBadge"),
            ("commercial-frontend/src/components/BookDetail.jsx",        "Nuevo",      "Ficha completa del libro"),
            ("commercial-frontend/src/components/CatalogFilters.jsx",    "Nuevo",      "Panel de filtros: categoría, precio, condición, editorial, año"),
            ("commercial-frontend/src/components/SearchBar.jsx",         "Nuevo",      "Barra de búsqueda con debounce (300 ms)"),
            ("commercial-frontend/src/components/AvailabilityBadge.jsx", "Nuevo",      "Badge de stock: disponible / agotado"),
            ("commercial-frontend/src/components/PriceBadge.jsx",        "Nuevo",      "Precio con indicador de fuente (eBay / Estimado / A consultar)"),
            ("commercial-frontend/src/components/EnrichedBookImage.jsx", "Nuevo",      "Imagen de portada con fallback a iniciales del título"),
            ("commercial-frontend/src/components/Cart.jsx",              "Nuevo",      "Carrito de compra básico"),
            ("commercial-frontend/src/api.js",                           "Nuevo",      "Funciones de llamada al BFF: books, categories, detail"),
            ("commercial-frontend/src/cart.js",                          "Nuevo",      "Estado del carrito con localStorage"),
            ("commercial-frontend/src/App.jsx",                          "Modificado", "Rutas: / → Catalogo, /book/:id → BookDetail"),
            ("commercial-frontend/src/index.css",                        "Modificado", "Estilos globales y utilidades Tailwind"),
        ],
        [3.5, 1.0, 3.5]
    )

    add_heading(doc, "Rutas del BFF que consume", 2, AZUL_MEDIO)
    add_table(doc,
        ["Método", "Ruta en BFF", "Para qué"],
        [
            ("GET", "/api/catalog/products",          "Listado paginado con filtros (?page=1&limit=20&category=...)"),
            ("GET", "/api/catalog/products/{id}",     "Detalle del producto con datos enriquecidos y precio"),
            ("GET", "/api/catalog/categories",     "Lista de categorías disponibles"),
        ],
        [0.8, 2.8, 4.4]
    )

    add_heading(doc, "Reglas de interfaz clave", 2, AZUL_MEDIO)
    add_table(doc,
        ["Situación", "Comportamiento"],
        [
            ("Sin portada disponible",            "Placeholder con iniciales del título (EnrichedBookImage)"),
            ("Sin precio calculado",              "Texto 'Precio a consultar' en lugar del PriceBadge"),
            ("Stock = 0",                         "Botón 'Agregar al carrito' deshabilitado; badge 'Agotado'"),
            ("is_fallback = true",                "PriceBadge muestra 'Estimado' en lugar del precio numérico"),
            ("Cargando",                          "Skeleton loader de 20 tarjetas mientras llegan datos"),
        ],
        [2.5, 5.5]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Levantar el proyecto completo:", bold=True)
    add_code(doc, "docker compose up --build")

    add_body(doc, "2. Abrir el catálogo en el navegador:", bold=True)
    add_code(doc, "http://localhost:3000")

    add_body(doc, "3. Puntos a verificar:", bold=True)
    for item in [
        "Catálogo carga con grid de libros y skeleton mientras espera",
        "Filtros por categoría y rango de precio funcionan",
        "SearchBar filtra resultados con debounce (escribe rápido → una sola petición)",
        "BookCard muestra precio si disponible, 'Precio a consultar' si no",
        "Clic en tarjeta abre BookDetail con descripción enriquecida",
        "Libros sin stock muestran badge 'Agotado' y botón deshabilitado",
        "EnrichedBookImage muestra placeholder con iniciales si no hay portada",
    ]:
        add_bullet(doc, item)

    add_body(doc, "4. Desarrollo local sin Docker:", bold=True)
    add_code(doc, """\
cd commercial-frontend
npm install
npm run dev
# → http://localhost:3000""")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Componente", "Qué verifica"],
        [
            ("BookCard renders price when available",           "BookCard",       "Precio numérico visible si está disponible"),
            ("BookCard shows 'Precio a consultar' when no price","BookCard",      "Texto correcto cuando no hay precio"),
            ("AvailabilityBadge shows correct state for stock=0","AvailabilityBadge","Badge 'Agotado' y botón deshabilitado"),
            ("SearchBar debounces input correctly",             "SearchBar",      "Solo una petición tras entrada rápida"),
            ("CatalogFilters apply and clear correctly",        "CatalogFilters", "Filtros modifican query params y se limpian"),
        ],
        [3.0, 1.5, 3.5]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 6 — Integración Real entre Microservicios (BFF)
# ────────────────────────────────────────────────────────────
def dev6(doc):
    add_heading(doc, "DEV 6 — Integración Real entre Microservicios", 1, AZUL_OSCURO)
    add_body(doc,
        "Responsable de que el BFF Gateway enrute correctamente hacia los servicios del Sprint 2 "
        "y que el health check consolidado refleje el estado real de todos los microservicios.",
        color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("bff-gateway/app/proxy.py",                    "Modificado", "Proxy genérico: forward_request() con timeout y manejo de errores"),
            ("bff-gateway/app/routers/gateway_router.py",   "Modificado", "Agrega rutas Sprint 2: /api/admin/pricing/*, /api/admin/enrichment/*"),
            ("bff-gateway/app/main.py",                     "Modificado", "Health check consolidado que agrega estado de todos los servicios"),
            ("bff-gateway/tests/test_gateway.py",           "Modificado", "Tests del BFF incluyendo nuevas rutas Sprint 2"),
        ],
        [3.2, 1.0, 3.8]
    )

    add_heading(doc, "Rutas añadidas al BFF para Sprint 2", 2, AZUL_MEDIO)
    add_table(doc,
        ["Ruta en BFF", "Proxea hacia"],
        [
            ("/api/admin/pricing",                           "pricing-service:8005/pricing"),
            ("/api/admin/pricing/{book_id}",                 "pricing-service:8005/pricing/{book_id}"),
            ("/api/admin/pricing/{book_id}/history",         "pricing-service:8005/pricing/{book_id}/history"),
            ("/api/admin/pricing/{book_id}/explanation",     "pricing-service:8005/pricing/{book_id}/explanation"),
            ("/api/admin/pricing/recalculate",               "pricing-service:8005/pricing/calculate (POST)"),
            ("/api/admin/enrichment/status",                 "ai-enrichment-service:8004/external-apis/status"),
            ("/api/admin/enrichment/trigger",                "ai-enrichment-service:8004/enrich (POST)"),
        ],
        [3.5, 4.5]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Health check consolidado:", bold=True)
    add_code(doc, "curl http://localhost:8009/health")
    add_body(doc, "Debe listar todos los servicios con 'status': 'ok'.", italic=True)

    add_body(doc, "2. Calcular precio a través del BFF:", bold=True)
    add_code(doc, """\
curl -X POST http://localhost:8009/api/admin/pricing/recalculate \\
  -H "Content-Type: application/json" \\
  -d '{"book_id": "...", "isbn": "...", "condition": "BUENO", "title": "..."}'""")

    add_body(doc, "3. Ver listado de precios vía BFF:", bold=True)
    add_code(doc, "curl http://localhost:8009/api/admin/pricing")

    add_body(doc, "4. Estado del enriquecimiento vía BFF:", bold=True)
    add_code(doc, "curl http://localhost:8009/api/admin/enrichment/status")

    add_body(doc, "5. Ejecutar tests del BFF:", bold=True)
    add_code(doc, """\
cd bff-gateway
pip install -r requirements.txt
pytest tests/ -v""")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Qué verifica"],
        [
            ("test_bff_health_aggregates_all_services",        "El BFF incluye estado de pricing y enrichment en /health"),
            ("test_pricing_proxy_route_returns_200",           "Ruta /api/admin/pricing proxea correctamente"),
            ("test_enrichment_trigger_proxy_works",            "POST /api/admin/enrichment/trigger llega al enrichment-service"),
            ("test_bff_handles_upstream_timeout",              "BFF devuelve 503 si servicio upstream no responde"),
        ],
        [3.5, 4.5]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 7 — Auditoría
# ────────────────────────────────────────────────────────────
def dev7(doc):
    add_heading(doc, "DEV 7 — Auditoría de Pricing e IA", 1, AZUL_OSCURO)
    add_body(doc,
        "Módulo de auditoría implementado dentro del BFF Gateway. Expone endpoints de consulta "
        "de decisiones de pricing y enriquecimiento, detecta anomalías (precios outlier, "
        "APIs que fallan repetidamente) y genera estadísticas por período.", color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("bff-gateway/DOCUMENTACION.md",             "Nuevo",      "Documentación interna del BFF y módulo de auditoría"),
            ("bff-gateway/app/routers/gateway_router.py","Modificado", "Agrega rutas /api/admin/audit/*"),
            ("bff-gateway/app/proxy.py",                 "Modificado", "Funciones forward para consultas de auditoría"),
        ],
        [3.2, 1.0, 3.8]
    )

    add_heading(doc, "Endpoints de auditoría", 2, AZUL_MEDIO)
    add_table(doc,
        ["Método", "Ruta", "Descripción"],
        [
            ("GET", "/api/admin/audit/pricing",            "Decisiones de pricing recientes"),
            ("GET", "/api/admin/audit/pricing/{id}",       "Detalle completo de una decisión"),
            ("GET", "/api/admin/audit/enrichment",         "Solicitudes de enriquecimiento recientes"),
            ("GET", "/api/admin/audit/enrichment/{id}",    "Detalle de un enriquecimiento"),
            ("GET", "/api/admin/audit/stats",              "Estadísticas generales (conteos, tasas de éxito)"),
            ("GET", "/api/admin/audit/anomalies",          "Precios outlier y alertas de APIs degradadas"),
        ],
        [0.8, 2.8, 4.4]
    )

    add_heading(doc, "Reglas de detección de anomalías", 2, AZUL_MEDIO)
    add_table(doc,
        ["Anomalía", "Condición"],
        [
            ("Precio outlier",         "Precio sugerido > 3× el promedio histórico del mismo libro"),
            ("API degradada",          "Fuente externa falla más de 5 veces en 1 hora"),
            ("Precio por debajo mínimo","Factor de condición produce precio < $1.00"),
        ],
        [2.5, 5.5]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Ver decisiones de pricing recientes:", bold=True)
    add_code(doc, "curl http://localhost:8009/api/admin/audit/pricing")

    add_body(doc, "2. Ver estadísticas del sistema:", bold=True)
    add_code(doc, "curl http://localhost:8009/api/admin/audit/stats")

    add_body(doc, "3. Ver anomalías detectadas:", bold=True)
    add_code(doc, "curl http://localhost:8009/api/admin/audit/anomalies")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Qué verifica"],
        [
            ("test_audit_entry_created_on_pricing_decision",     "Cada decisión de pricing genera entrada de auditoría"),
            ("test_audit_entry_created_on_enrichment_result",    "Cada enriquecimiento completado genera entrada"),
            ("test_outlier_detection_flags_price_3x_average",    "Precio outlier (>3× promedio) es detectado"),
            ("test_stats_endpoint_returns_correct_counts",       "Endpoint /stats devuelve conteos correctos"),
            ("test_audit_log_includes_all_decision_factors",     "Registro incluye todos los factores de la decisión"),
        ],
        [3.5, 4.5]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 8 — APIs Externas Robustez
# ────────────────────────────────────────────────────────────
def dev8(doc):
    add_heading(doc, "DEV 8 — Integración con APIs Externas (Robustez)", 1, AZUL_OSCURO)
    add_body(doc,
        "Responsable de la robustez de todas las integraciones externas del Sprint 2. "
        "Implementa el patrón base de circuit breaker y caché en memoria compartido por "
        "Google Books, Open Library, Crossref y eBay.", color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("ai-enrichment-service/app/infrastructure/adapters/base_adapter.py", "Nuevo", "SimpleCache (TTL dict) + CircuitBreaker (5 fallos, 5 min cooldown)"),
            ("ai-enrichment-service/app/infrastructure/adapters/google_books_adapter.py","Nuevo","Hereda base_adapter; cache TTL=3600; MAX_RETRIES=3; TIMEOUT=5s"),
            ("ai-enrichment-service/app/infrastructure/adapters/open_library_adapter.py","Nuevo","Hereda base_adapter; búsqueda por ISBN y query"),
            ("ai-enrichment-service/app/infrastructure/adapters/crossref_adapter.py",    "Nuevo","Hereda base_adapter; búsqueda de works académicos"),
            ("pricing-service/app/infrastructure/adapters/ebay_adapter.py",              "Nuevo","Circuit breaker + caché para eBay Browse API"),
        ],
        [3.5, 1.0, 3.5]
    )

    add_heading(doc, "Parámetros del circuit breaker", 2, AZUL_MEDIO)
    add_table(doc,
        ["Parámetro", "Valor", "Descripción"],
        [
            ("MAX_FAILURES",     "5",    "Fallos consecutivos antes de abrir el circuito"),
            ("COOLDOWN_SECONDS", "300",  "Segundos de espera antes de intentar de nuevo (5 min)"),
            ("Estado OPEN",      "—",    "La API se marca como degradada; se retorna None directamente"),
            ("Estado CLOSED",    "—",    "La API funciona con normalidad"),
            ("Estado HALF-OPEN", "—",    "Después del cooldown; primer intento decide si cierra o reabre"),
        ],
        [2.0, 1.0, 5.0]
    )

    add_heading(doc, "Parámetros del caché en memoria", 2, AZUL_MEDIO)
    add_table(doc,
        ["Parámetro", "Valor", "Descripción"],
        [
            ("CACHE_TTL",     "3600 s", "Una hora de validez por entrada"),
            ("Clave",         "isbn o query", "Normalizada a minúsculas"),
            ("Persistencia",  "Ninguna", "El caché se vacía al reiniciar el contenedor (aceptable en Sprint 2)"),
        ],
        [1.5, 1.5, 5.0]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Verificar circuit breaker activo (eBay):", bold=True)
    add_code(doc, "curl http://localhost:8005/pricing/external-apis/status")

    add_body(doc, "2. Verificar circuit breakers de enrichment:", bold=True)
    add_code(doc, "curl http://localhost:8004/external-apis/status")
    add_body(doc, "Respuesta esperada:", italic=True)
    add_code(doc, """\
{
  "google_books":  {"status": "ok", "failures": 0},
  "open_library":  {"status": "ok", "failures": 0},
  "crossref":      {"status": "ok", "failures": 0}
}""")

    add_body(doc, "3. Probar caché: ejecutar la misma búsqueda dos veces y verificar logs:", bold=True)
    add_code(doc, """\
# Primera llamada → llama a la API externa
curl -X POST http://localhost:8004/enrich -H "Content-Type: application/json" \\
  -d '{"book_id":"test-1","isbn":"9780141439518","title":"Great Expectations"}'

# Segunda llamada → debería venir del caché (verificar logs del contenedor)
curl -X POST http://localhost:8004/enrich -H "Content-Type: application/json" \\
  -d '{"book_id":"test-2","isbn":"9780141439518","title":"Great Expectations"}'

# Ver logs del servicio
docker compose logs ai-enrichment-service --tail=20""")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Qué verifica"],
        [
            ("test_google_books_cache_hit_avoids_http_call",   "Segunda consulta igual no hace petición HTTP"),
            ("test_open_library_returns_none_on_404",          "404 retorna None sin lanzar excepción"),
            ("test_circuit_breaker_opens_after_5_failures",    "Circuito se abre tras 5 fallos consecutivos"),
            ("test_circuit_breaker_resets_after_cooldown",     "Circuito se cierra tras pasar el cooldown"),
            ("test_crossref_handles_rate_limit_429",           "Rate limit (429) maneja graciosamente sin crash"),
        ],
        [3.5, 4.5]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# DEV 9 — Optimización y Persistencia
# ────────────────────────────────────────────────────────────
def dev9(doc):
    add_heading(doc, "DEV 9 — Optimización y Persistencia", 1, AZUL_OSCURO)
    add_body(doc,
        "Responsable de la orquestación Docker del Sprint 2, los esquemas de base de datos "
        "de los nuevos servicios y los datos semilla para el equipo de desarrollo.",
        color=GRIS_TEXTO)
    add_separator(doc)

    add_heading(doc, "Archivos trabajados", 2, AZUL_MEDIO)
    add_table(doc,
        ["Archivo", "Tipo", "Descripción"],
        [
            ("docker-compose.yml",           "Modificado", "Agrega pricing-service, ai-enrichment-service, pricing-db, enrichment-db"),
            (".env / .env.example",          "Modificado", "Agrega DB_NAME_ENRICHMENT, DB_NAME_PRICING, GOOGLE_BOOKS_API_KEY, EBAY_APP_ID"),
            ("ai-enrichment-service/app/infrastructure/database/models.py",  "Nuevo", "Schema: enrichment_requests, enrichment_results"),
            ("pricing-service/app/infrastructure/database/models.py",         "Nuevo", "Schema: price_references, pricing_decisions"),
            ("ai-enrichment-service/app/infrastructure/database/connection.py","Nuevo","Creación async de tablas en lifespan"),
            ("pricing-service/app/infrastructure/database/connection.py",      "Nuevo","Creación async de tablas en lifespan"),
            ("README.md",                    "Modificado", "Instrucciones de setup para Sprint 1 + Sprint 2"),
        ],
        [3.2, 1.0, 3.8]
    )

    add_heading(doc, "Servicios Docker añadidos en Sprint 2", 2, AZUL_MEDIO)
    add_table(doc,
        ["Servicio", "Puerto", "Imagen / Build", "Depende de"],
        [
            ("pricing-service",        "8005", "build: ./pricing-service",       "pricing-db, catalog-service, inventory-service"),
            ("ai-enrichment-service",  "8004", "build: ./ai-enrichment-service", "enrichment-db, catalog-service"),
            ("pricing-db",             "—",    "postgres:15",                    "—"),
            ("enrichment-db",          "—",    "postgres:15",                    "—"),
        ],
        [2.0, 0.8, 2.0, 3.2]
    )

    add_heading(doc, "Esquema de tablas de BD nuevas", 2, AZUL_MEDIO)
    add_body(doc, "pricing_db:", bold=True)
    add_table(doc,
        ["Tabla", "Columnas clave"],
        [
            ("price_references",  "id, source, external_price, currency, observed_at, book_reference, confidence_score"),
            ("pricing_decisions", "id, book_reference, suggested_price, explanation, condition_factor, reference_count, base_price, is_fallback, created_at"),
        ],
        [2.0, 6.0]
    )
    add_body(doc, "enrichment_db:", bold=True)
    add_table(doc,
        ["Tabla", "Columnas clave"],
        [
            ("enrichment_requests", "id, book_id, isbn, title, author, status, source_used, error_message, requested_at"),
            ("enrichment_results",  "id, request_id, normalized_title, normalized_author, cover_url, confidence_score, created_at"),
        ],
        [2.0, 6.0]
    )

    add_heading(doc, "Cómo probar", 2, AZUL_MEDIO)
    add_body(doc, "1. Primer arranque completo:", bold=True)
    add_code(doc, """\
# Limpiar volúmenes previos (OBLIGATORIO si ya corrió Sprint 1)
docker compose down -v --rmi all

# Construir y levantar todo (incluyendo Sprint 2)
docker compose up --build""")

    add_body(doc, "2. Verificar que los contenedores nuevos están corriendo:", bold=True)
    add_code(doc, """\
docker compose ps
# Debe mostrar: pricing-service, ai-enrichment-service, pricing-db, enrichment-db""")

    add_body(doc, "3. Verificar health checks de los servicios nuevos:", bold=True)
    add_code(doc, """\
curl http://localhost:8004/health   # AI Enrichment Service
curl http://localhost:8005/health   # Pricing Service""")

    add_body(doc, "4. Verificar health check consolidado del BFF:", bold=True)
    add_code(doc, """\
curl http://localhost:8009/health
# Debe incluir "pricing" y "enrichment-real" con status "ok" """)

    add_body(doc, "5. Arranques posteriores (sin rebuild):", bold=True)
    add_code(doc, "docker compose up")

    add_heading(doc, "Tests obligatorios", 2, AZUL_MEDIO)
    add_table(doc,
        ["Test", "Qué verifica"],
        [
            ("test_all_services_start_in_docker_compose",   "Todos los contenedores pasan a estado 'healthy'"),
            ("test_migrations_run_successfully_on_clean_db","Tablas nuevas se crean sin errores en BD limpia"),
            ("test_seed_data_loads_without_errors",          "Scripts de semilla no generan errores ni conflictos"),
            ("test_pricing_query_uses_index",                "Query de historial usa índice en created_at DESC"),
        ],
        [3.5, 4.5]
    )
    doc.add_page_break()

# ────────────────────────────────────────────────────────────
# SECCIÓN FINAL: comandos generales y contacto
# ────────────────────────────────────────────────────────────
def closing_section(doc):
    add_heading(doc, "Referencia Rápida — Todos los Servicios", 1, AZUL_OSCURO)
    add_separator(doc)

    add_heading(doc, "Puertos del sistema", 2, AZUL_MEDIO)
    add_table(doc,
        ["Servicio", "Puerto", "URL Swagger/Docs"],
        [
            ("commercial-frontend",    "3000", "http://localhost:3000"),
            ("admin-frontend",         "3001", "http://localhost:3001"),
            ("auth-service",           "8001", "http://localhost:8001/docs"),
            ("inventory-service",      "8002", "http://localhost:8002/docs"),
            ("catalog-service",        "8003", "http://localhost:8003/docs"),
            ("ai-enrichment-service",  "8004", "http://localhost:8004/docs  ← SPRINT 2"),
            ("pricing-service",        "8005", "http://localhost:8005/docs  ← SPRINT 2"),
            ("ai-enrichment-mock",     "8006", "http://localhost:8006/docs"),
            ("data-quality-module",    "8007", "http://localhost:8007/docs"),
            ("config-module",          "8008", "http://localhost:8008/docs"),
            ("bff-gateway",            "8009", "http://localhost:8009/docs"),
        ],
        [2.0, 1.0, 5.0]
    )

    add_heading(doc, "Comandos esenciales", 2, AZUL_MEDIO)
    add_table(doc,
        ["Comando", "Para qué"],
        [
            ("docker compose up --build",                   "Primer arranque o rebuild completo"),
            ("docker compose up",                           "Arranque normal (sin rebuild)"),
            ("docker compose down -v --rmi all",            "Limpiar completamente (volúmenes + imágenes)"),
            ("docker compose logs <servicio> -f",           "Ver logs en tiempo real de un servicio"),
            ("docker compose ps",                           "Ver estado de todos los contenedores"),
            ("pytest tests/ -v --cov=app",                  "Correr tests con cobertura (desde carpeta del servicio)"),
            ("curl http://localhost:8009/health",           "Verificar estado de todos los servicios"),
        ],
        [3.5, 4.5]
    )

    add_heading(doc, "Variables de entorno opcionales (Sprint 2)", 2, AZUL_MEDIO)
    add_body(doc,
        "Los servicios funcionan sin estas keys usando fallbacks internos. Agregarlas mejora la "
        "calidad de los resultados.", color=GRIS_TEXTO)
    add_table(doc,
        ["Variable", "Servicio", "Sin key"],
        [
            ("GOOGLE_BOOKS_API_KEY", "ai-enrichment-service", "Usa Open Library + Crossref como fallback"),
            ("EBAY_APP_ID",          "pricing-service",        "Usa reglas internas de pricing"),
        ],
        [2.2, 2.0, 3.8]
    )

    add_separator(doc)
    add_body(doc, "BookFlow AI Commerce — Sprint 2  ·  Abril 2026  ·  Equipo de Desarrollo",
             color=GRIS_TEXTO, italic=True)

# ────────────────────────────────────────────────────────────
# MAIN
# ────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    cover_page(doc)
    index_page(doc)
    dev1(doc)
    dev2(doc)
    dev3(doc)
    dev4(doc)
    dev5(doc)
    dev6(doc)
    dev7(doc)
    dev8(doc)
    dev9(doc)
    closing_section(doc)

    output = "c:/Users/jandr/OneDrive/Desktop/BookFlow_Sprint2_Documentacion_Devs.docx"
    doc.save(output)
    print(f"Documento generado: {output}")

if __name__ == "__main__":
    main()
